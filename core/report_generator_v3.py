"""学情报告生成器V3 — 精简分析 + 正式报告 + Word文件生成"""
import os
import sys
import re
import hashlib
import json
import pandas as pd
from openai import OpenAI
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# 注册中文字体（Windows + Linux 双环境兼容）
_FONT_REGISTERED = False
def _register_chinese_font():
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return

    _has_normal = False
    _has_bold = False

    font_paths = [
        # === Windows ===
        (r"C:\Windows\Fonts\msyh.ttc", "MSYH"),
        (r"C:\Windows\Fonts\msyhbd.ttc", "MSYHBD"),
        (r"C:\Windows\Fonts\simsun.ttc", "SimSun"),
        (r"C:\Windows\Fonts\simhei.ttf", "SimHei"),
        # === Linux (Streamlit Cloud / Docker) ===
        # Noto Sans CJK (fonts-noto-cjk)
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "MSYH"),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc", "MSYHBD"),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc", "MSYH"),
        ("/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc", "MSYHBD"),
        # DroidSansFallback（Debian 默认中文字体，无需额外安装）
        ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf", "MSYH"),
        ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf", "MSYHBD"),
        # WenQuanYi Micro Hei（轻量中文字体）
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "MSYH"),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", "MSYHBD"),
    ]
    for path, name in font_paths:
        try:
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont(name, path))
                if name == "MSYH":
                    _has_normal = True
                elif name == "MSYHBD":
                    _has_bold = True
        except:
            pass

    # 设置字体族映射：告诉 reportlab MSYHBD 是 MSYH 的粗体
    # 这是修复 "Can't map determine family/bold/italic for mshbd" 的关键
    if _has_normal or _has_bold:
        normal = "MSYH" if _has_normal else "MSYHBD"
        bold = "MSYHBD" if _has_bold else "MSYH"
        try:
            pdfmetrics.registerFontFamily("MSYH", normal=normal, bold=bold)
        except:
            pass

    # 最终兜底：如果字体都没注册成功，使用 Helvetica 避免崩溃
    if not _has_normal and not _has_bold:
        print("[PDF] 警告：未找到任何中文字体，PDF 中文将无法显示")
        try:
            pdfmetrics.registerFontFamily("MSYH", normal="Helvetica", bold="Helvetica-Bold")
        except:
            pass

    _FONT_REGISTERED = True

_register_chinese_font()

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import OPENAI_API_KEY, OPENAI_BASE_URL, MODEL_NAME, TEMPERATURE, MAX_TOKENS
from core.prompt_utils import append_constraints

client = OpenAI(api_key=OPENAI_API_KEY, base_url=OPENAI_BASE_URL)
MAX_WORKERS = 20

# ===== LLM响应缓存（文件级，避免重复生成）=====
_CACHE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "cache", "llm_responses")
os.makedirs(_CACHE_DIR, exist_ok=True)

def _cache_key(prompt: str, system_msg: str) -> str:
    """生成唯一缓存key"""
    raw = f"{prompt}|||{system_msg}"
    return hashlib.md5(raw.encode("utf-8")).hexdigest()

def _get_cached(key: str) -> str:
    """从缓存读取"""
    cache_path = os.path.join(_CACHE_DIR, f"{key}.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("response", "")
        except:
            pass
    return ""

def _set_cache(key: str, response: str):
    """写入缓存"""
    cache_path = os.path.join(_CACHE_DIR, f"{key}.json")
    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump({"response": response}, f, ensure_ascii=False)
    except:
        pass


def load_prompt(template_name: str) -> str:
    prompt_dir = os.path.join(os.path.dirname(__file__), "prompts")
    filepath = os.path.join(prompt_dir, template_name)
    with open(filepath, "r", encoding="utf-8") as f:
        return append_constraints(f.read())


def build_student_data_text(row_data: dict) -> str:
    lines = []
    for key, val in row_data.items():
        if pd.notna(val) and str(val).strip():
            val_str = str(val)
            if len(val_str) > 100:
                val_str = val_str[:100] + "..."
            lines.append(f"- {key}: {val_str}")
    return "\n".join(lines)


def _calc_column_data_rate(all_students_data: list, col_name: str) -> float:
    """计算某列有数据的比例"""
    count = 0
    for item in all_students_data:
        val = item["row_data"].get(col_name)
        if pd.notna(val) and str(val).strip() and str(val).strip() not in ["", "无", "-", "N/A", "n/a", "0", "0%"]:
            count += 1
    return count / len(all_students_data) if all_students_data else 0


def classify_student(row_data: dict, all_students_data: list = None,
                     lectures: list = None, attendance_rates: dict = None) -> str:
    """
    根据数据快速分类：优秀/中等/差
    优先使用讲次数据（lectures + attendance_rates），
    若没有则回退到原始列名匹配
    """
    if lectures and attendance_rates:
        # 使用讲次数据分析
        score = 0
        valid_lecture_count = 0
        for lec in lectures:
            lec_name = lec["lecture"]
            rate = attendance_rates.get(lec_name, 0)
            if rate < 0.1:
                continue  # 未开课，跳过
            valid_lecture_count += 1
            cols = lec["cols"]

            # 是否有效听课
            eff_col = cols.get("是否有效听课")
            if eff_col and eff_col in row_data:
                v = str(row_data[eff_col]).strip()
                if v in ["是", "1", "True", "true", "有效"]:
                    score += 2
                else:
                    score -= 2

            # 答题正确率
            acc_col = cols.get("直播答题正确率")
            if acc_col and acc_col in row_data:
                try:
                    pct = float(re.search(r'(\d+(?:\.\d+)?)', str(row_data[acc_col])).group(1))
                    if pct >= 80:
                        score += 2
                    elif pct <= 70:
                        score -= 2
                except:
                    pass

            # 练习得分
            score_col = cols.get("练习得分")
            if score_col and score_col in row_data:
                try:
                    s = float(re.search(r'(\d+(?:\.\d+)?)', str(row_data[score_col])).group(1))
                    if s >= 80:
                        score += 2
                    elif s <= 70:
                        score -= 1
                except:
                    pass

        if valid_lecture_count == 0:
            # 讲次数据无法判断，尝试原始列名匹配
            return _classify_by_raw_columns(row_data, all_students_data)

        if score >= 5:
            return "优秀"
        elif score >= -2:
            return "中等"
        else:
            return "差"

    return _classify_by_raw_columns(row_data, all_students_data)


def _classify_by_raw_columns(row_data: dict, all_students_data: list = None) -> str:
    """原始列名匹配分类（回退方案）"""
    score = 0
    total_weight = 0
    
    for key, val in row_data.items():
        if pd.isna(val):
            continue
        k = str(key)
        v = str(val).strip()
        if not v or v in ["无", "-", "N/A", "n/a"]:
            continue
        
        if all_students_data:
            data_rate = _calc_column_data_rate(all_students_data, key)
            if data_rate < 0.30:
                continue
        
        total_weight += 1
        
        if "到课" in k or "出勤" in k:
            if v in ["到课", "已到", "全勤", "100%"]:
                score += 2
            elif v in ["缺课", "未到", "旷课"]:
                score -= 2
        if "作业完成" in k or "作业" in k:
            try:
                pct = float(re.search(r'(\d+)', v).group(1))
                if pct >= 90: score += 2
                elif pct >= 70: score += 1
                elif pct < 50: score -= 2
            except:
                pass
        if "正确率" in k or "正确" in k:
            try:
                pct = float(re.search(r'(\d+)', v).group(1))
                if pct >= 80: score += 1
                elif pct < 50: score -= 1
            except:
                pass
        if "成绩" in k or "分数" in k or "测验" in k:
            try:
                s = float(re.search(r'(\d+(?:\.\d+)?)', v).group(1))
                if s >= 80: score += 2
                elif s >= 60: score += 1
                elif s < 40: score -= 2
            except:
                pass
    
    if total_weight == 0:
        return "中等"
    
    if score >= 4:
        return "优秀"
    elif score >= 0:
        return "中等"
    else:
        return "差"


def _call_llm(prompt: str, max_retry: int = 2) -> str:
    """调用LLM，带缓存和重试机制"""
    system_msg = "你是专业英语学情报告单撰写助手。严格依据原始数据生成报告，禁止编造。"
    # 检查缓存
    ckey = _cache_key(prompt, system_msg)
    cached = _get_cached(ckey)
    if cached:
        print(f"[LLM] 命中缓存")
        return cached

    last_error = ""
    for attempt in range(max_retry):
        try:
            response = client.chat.completions.create(
                model=MODEL_NAME,
                temperature=0.3,
                max_tokens=1500,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ]
            )
            content = response.choices[0].message.content
            if content and content.strip():
                _set_cache(ckey, content)
                return content
            last_error = "LLM返回空内容"
            print(f"[LLM] 空返回，重试 {attempt+1}/{max_retry}")
        except Exception as e:
            last_error = str(e)
            print(f"[LLM] 异常: {e}，重试 {attempt+1}/{max_retry}")
            if attempt < max_retry - 1:
                import time as _time
                _time.sleep(0.5)
    return f"报告生成失败，请重新生成\n\n错误详情: {last_error}"


def generate_single_report(item: dict, lectures: list = None, attendance_rates: dict = None) -> dict:
    """为单个学生生成报告"""
    try:
        template = load_prompt("batch_report_v3.txt")
        student_data = build_student_data_text(item["row_data"])
        
        # 讲次分析
        lecture_context = ""
        if lectures and attendance_rates:
            from core.lecture_parser import analyze_student_lectures, build_lecture_context as build_ctx
            lec_analysis = analyze_student_lectures(item["row_data"], lectures, attendance_rates)
            lecture_context = build_ctx(lec_analysis, lectures)
        
        # 安全格式化：转义 student_data 中的花括号，防止 format() 报错
        safe_student_data = student_data.replace("{", "{{").replace("}", "}}")
        safe_lecture_context = lecture_context.replace("{", "{{").replace("}", "}}")
        
        prompt = template.format(
            name=item["name"],
            lecture_context=safe_lecture_context,
            student_data=safe_student_data
        )
        
        report_text = _call_llm(prompt)
        
        # 整个输出就是报告（不再分割快速分析和正式报告）
        analysis = ""
        formal_report = report_text
        
        return {
            "学生姓名": item["name"],
            "学生ID": item.get("student_id", ""),
            "分类": item.get("category", ""),
            "快速分析": analysis,
            "正式报告": formal_report,
            "完整输出": report_text,
        }
    except Exception as e:
        return {
            "学生姓名": item.get("name", "未知"),
            "学生ID": item.get("student_id", ""),
            "分类": item.get("category", ""),
            "快速分析": "",
            "正式报告": f"报告生成异常: {e}",
            "完整输出": f"报告生成异常: {e}",
        }


def batch_generate_reports(students_data: list, progress_callback=None, lectures: list = None, attendance_rates: dict = None) -> list:
    """批量生成报告"""
    results = [None] * len(students_data)
    completed = 0
    total = len(students_data)

    def _task(idx, item):
        try:
            return idx, generate_single_report(item, lectures, attendance_rates)
        except Exception as e:
            return idx, {
                "学生姓名": item["name"],
                "学生ID": item.get("student_id", ""),
                "分类": item["category"],
                "快速分析": f"生成失败: {e}",
                "正式报告": "",
                "完整输出": "",
            }

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(_task, i, s): i for i, s in enumerate(students_data)}
        for future in as_completed(futures):
            idx, result = future.result()
            results[idx] = result
            completed += 1
            if progress_callback:
                progress_callback(completed, total)

    return [r for r in results if r is not None]


def generate_word_report(student: dict, output_dir: str) -> str:
    """为单个学生生成正式Word报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 从报告内容提取标题（第一行 # 开头）
    report = student.get("正式报告", student.get("完整输出", ""))
    title_text = f"{student['学生姓名']} · 英语学情报告单"
    for line in report.split("\n"):
        line = line.strip()
        if line.startswith("#"):
            title_text = line.lstrip("# ").strip()
            break

    # 标题
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    
    doc.add_paragraph()  # 空行
    
    # 解析报告内容
    sections = {
        "学情总览": [], "答题表现": [], "学习亮点": [],
        "待优化": [], "学习安排": [], "寄语": []
    }
    
    current_section = None
    for line in report.split("\n"):
        line = line.strip()
        if not line or line == "---" or line.startswith("#"):
            continue
        matched = False
        for section_name in sections:
            if section_name in line:
                current_section = section_name
                matched = True
                break
        if not matched and current_section:
            sections[current_section].append(line)
    
    # 写入各部分
    section_titles = {
        "学情总览": "📊 学情总览",
        "答题表现": "📈 答题表现",
        "学习亮点": "✨ 学习亮点",
        "待优化": "⚠️ 待优化事项",
        "学习安排": "📋 学习安排",
        "寄语": "💌 教师简短寄语",
    }
    
    for section_name, lines in sections.items():
        if lines:
            h = doc.add_heading(section_titles.get(section_name, section_name), level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                run.font.size = Pt(13)
            
            for line in lines:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(4)
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—— 高途教育 ——")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(9)
    
    # 保存
    safe_name = re.sub(r'[\\/:*?"<>|]', '', student["学生姓名"])
    filename = f"学情报告_{safe_name}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def generate_category_word_report(category: str, students: list, output_dir: str) -> str:
    """生成某分类的批量汇总Word报告"""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    category_labels = {"优秀": "🌟 优秀学员", "中等": "📊 中等学员", "差": "⚠️ 需关注学员"}
    
    title = doc.add_heading(f'高途教育{category_labels.get(category, category)}学情汇总报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    
    count = doc.add_paragraph(f"共 {len(students)} 名学员")
    count.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for student in students:
        h = doc.add_heading(f"【{student['学生姓名']}】", level=2)
        report = student.get("正式报告", student.get("快速分析", ""))
        for line in report.split("\n"):
            line = line.strip()
            if line and not line.startswith("##") and not line.startswith("#"):
                doc.add_paragraph(line)
        doc.add_paragraph()  # 分隔
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—— 高途教育 ——")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(9)
    
    filename = f"学情汇总_{category}学员.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def _markdown_to_reportlab(text: str) -> str:
    """将简单markdown格式转换为reportlab支持的HTML标签格式"""
    # 先转markdown格式为HTML标签，再转义文本中的XML特殊字符
    # 使用占位符保护HTML标签
    bold_placeholders = {}
    def _replace_bold(m):
        idx = len(bold_placeholders)
        ph = f"__BOLD_{idx}__"
        bold_placeholders[ph] = f"<b>{m.group(1)}</b>"
        return ph
    text = re.sub(r'\*\*(.+?)\*\*', _replace_bold, text)
    # 标题
    text = re.sub(r'^#\s+(.+)$', _replace_bold, text, flags=re.MULTILINE)
    # 转义XML特殊字符（仅文本部分）
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    # 恢复占位符
    for ph, tag in bold_placeholders.items():
        text = text.replace(ph, tag)
    return text


def generate_pdf_report(student: dict, output_dir: str) -> str:
    """为单个学生生成正式PDF报告"""
    safe_name = re.sub(r'[\\/:*?"<>|]', '', student["学生姓名"])
    filename = f"学情报告_{safe_name}.pdf"
    filepath = os.path.join(output_dir, filename)

    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()

    # 自定义样式（使用中文字体）
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Title'],
        fontName='MSYHBD', fontSize=18,
        textColor=HexColor('#1A56DB'), spaceAfter=6
    )
    section_style = ParagraphStyle(
        'Section', parent=styles['Normal'],
        fontName='MSYHBD', fontSize=12,
        textColor=HexColor('#1A56DB'), spaceBefore=10, spaceAfter=4
    )
    body_style = ParagraphStyle(
        'Body', parent=styles['Normal'],
        fontName='MSYH', fontSize=10.5,
        textColor=HexColor('#333333'), spaceAfter=4,
        leading=18
    )
    body_bold_style = ParagraphStyle(
        'BodyBold', parent=styles['Normal'],
        fontName='MSYHBD', fontSize=10.5,
        textColor=HexColor('#333333'), spaceAfter=4,
        leading=18
    )
    list_style = ParagraphStyle(
        'List', parent=styles['Normal'],
        fontName='MSYH', fontSize=10.5,
        textColor=HexColor('#333333'), spaceAfter=3,
        leading=18, leftIndent=12
    )
    quote_style = ParagraphStyle(
        'Quote', parent=styles['Normal'],
        fontName='MSYH', fontSize=10.5,
        textColor=HexColor('#555555'), spaceAfter=6,
        leading=18, leftIndent=6,
        borderColor=HexColor('#1A56DB'), borderWidth=0,
        borderPadding=0
    )

    # 从报告内容提取标题
    report = student.get("正式报告", student.get("完整输出", ""))
    title_text = f"{student['学生姓名']} · 英语学情报告单"
    for line in report.split("\n"):
        line = line.strip()
        if line.startswith("#"):
            title_text = line.lstrip("# ").strip()
            break

    elements = []

    # 标题
    elements.append(Paragraph(title_text, title_style))
    elements.append(Spacer(1, 6))

    # 解析报告内容 - 按模块分割
    section_headers = ["📊 **学情总览**", "📈 **答题表现**", "✨ **学习亮点**",
                       "⚠️ **待优化事项**", "📋 **学习提升规划**", "💌 **教师寄语**"]
    # 也支持无emoji版本
    section_headers_fallback = ["学情总览", "答题表现", "学习亮点",
                                "待优化事项", "学习提升规划", "教师寄语"]

    # 按行解析
    current_section = None
    section_content = {}  # section_name -> [lines]

    for line in report.split("\n"):
        line_stripped = line.strip()
        if not line_stripped or line_stripped == "---":
            continue
        if line_stripped.startswith("#") and "英语学情报告单" in line_stripped:
            continue

        # 检测是否为章节标题
        matched = False
        for i, header in enumerate(section_headers):
            # 清理markdown符号后匹配
            clean_header = header.replace("**", "")
            if clean_header in line_stripped or line_stripped == clean_header:
                current_section = section_headers_fallback[i]
                if current_section not in section_content:
                    section_content[current_section] = []
                matched = True
                break
        if matched:
            continue

        # 后备匹配（无emoji）
        if not matched:
            for i, fallback in enumerate(section_headers_fallback):
                if f"**{fallback}**" in line_stripped or line_stripped == f"**{fallback}**" or line_stripped == fallback:
                    current_section = fallback
                    if current_section not in section_content:
                        section_content[current_section] = []
                    matched = True
                    break
        if matched:
            continue

        if current_section and line_stripped:
            if current_section not in section_content:
                section_content[current_section] = []
            section_content[current_section].append(line_stripped)

    # 按固定顺序输出
    section_order = ["学情总览", "答题表现", "学习亮点", "待优化事项", "学习提升规划", "教师寄语"]
    section_labels = {
        "学情总览": "📊 学情总览",
        "答题表现": "📈 答题表现",
        "学习亮点": "✨ 学习亮点",
        "待优化事项": "⚠️ 待优化事项",
        "学习提升规划": "📋 学习提升规划",
        "教师寄语": "💌 教师寄语",
    }

    for sec_name in section_order:
        lines = section_content.get(sec_name, [])
        if not lines:
            # 如果LLM没输出该模块，尝试从原始报告中提取
            continue

        elements.append(Paragraph(section_labels.get(sec_name, sec_name), section_style))

        for line in lines:
            if not line:
                continue
            # 检测是否为列表项（• 或 ①②③ 开头）
            is_list = line.startswith("•") or line.startswith("①") or line.startswith("②") or line.startswith("③")
            # 检测是否为带冒号/｜的横向数据行
            is_data_row = "｜" in line or "|" in line

            # 转换markdown为reportlab格式
            formatted = _markdown_to_reportlab(line)

            if sec_name == "教师寄语":
                if "高途辅导老师" in line:
                    elements.append(Paragraph(formatted, quote_style))
                else:
                    elements.append(Paragraph(formatted, quote_style))
            elif is_list:
                elements.append(Paragraph(formatted, list_style))
            elif is_data_row:
                elements.append(Paragraph(formatted, body_bold_style))
            else:
                elements.append(Paragraph(formatted, body_style))

    # 底部footer
    elements.append(Spacer(1, 20))
    elements.append(Paragraph("—— 高途教育 ——", ParagraphStyle(
        'Footer', parent=styles['Normal'],
        fontName='MSYH', fontSize=9,
        textColor=HexColor('#999999'), alignment=1
    )))

    doc.build(elements)
    return filepath


def generate_category_pdf_report(category: str, students: list, output_dir: str) -> str:
    """生成某分类的批量汇总PDF报告"""
    safe_cat = re.sub(r'[\\/:*?"<>|]', '', category)
    filename = f"学情汇总_{safe_cat}学员.pdf"
    filepath = os.path.join(output_dir, filename)

    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'Title', parent=styles['Title'],
        fontName='MSYHBD', fontSize=16,
        textColor=HexColor('#1A56DB'), spaceAfter=6
    )
    name_style = ParagraphStyle(
        'Name', parent=styles['Normal'],
        fontName='MSYHBD', fontSize=12,
        textColor=HexColor('#333333'), spaceBefore=12, spaceAfter=4
    )
    body_style = ParagraphStyle(
        'Body', parent=styles['Normal'],
        fontName='MSYH', fontSize=10,
        textColor=HexColor('#333333'), spaceAfter=3,
        leading=16
    )
    footer_style = ParagraphStyle(
        'Footer', parent=styles['Normal'],
        fontName='MSYH', fontSize=9,
        textColor=HexColor('#999999'), alignment=1
    )

    elements = []
    category_labels = {"优秀": "优秀学员", "中等": "中等学员", "差": "需关注学员"}
    elements.append(Paragraph(f"高途教育{category_labels.get(category, category)}学情汇总", title_style))
    elements.append(Paragraph(f"共 {len(students)} 名学员", body_style))

    for student in students:
        elements.append(Paragraph(f"【{student['学生姓名']}】", name_style))
        report = student.get("正式报告", student.get("快速分析", ""))
        for line in report.split("\n"):
            line = line.strip()
            if line and not line.startswith("##") and not line.startswith("#"):
                clean = line.replace("**", "")
                if clean:
                    elements.append(Paragraph(clean, body_style))

    elements.append(Spacer(1, 20))
    elements.append(Paragraph("—— 高途教育 ——", footer_style))

    doc.build(elements)
    return filepath
